"""
Bygg metodedokument for den feltspesifikke realiserte oljepris-modellen.
python-docx (node ikke installert).
"""
import shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path("/Users/emmastrandenskaar/Documents/Claude/Projects/Oljepris")
PROC = ROOT / "data" / "processed"
OUT = ROOT / "docs" / "NCS_Realized_Price_Methodology.docx"
ONEDRIVE = Path("/Users/emmastrandenskaar/Library/CloudStorage/OneDrive-BINorwegianBusinessSchool(BIEDU)/NCS_Realized_Price_Methodology.docx")

DARK = RGBColor(0x1F, 0x4E, 0x79); GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28); ORANGE = RGBColor(0xE6, 0x51, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59); LGREY = RGBColor(0x80, 0x80, 0x80); WHITE = RGBColor(0xFF, 0xFF, 0xFF)
GREEN_BG = "E8F5E9"; BLUE_BG = "E3F2FD"; ORANGE_BG = "FFF3E0"; HDR = "1F4E79"

doc = Document()
s = doc.sections[0]
s.page_height, s.page_width = Inches(11), Inches(8.5)
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Inches(1)
doc.styles['Normal'].font.name = 'Arial'; doc.styles['Normal'].font.size = Pt(11)

def shade(cell, color):
    tcpr = cell._tc.get_or_add_tcPr(); sh = OxmlElement('w:shd')
    sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear'); tcpr.append(sh)
def borders(cell, color="CCCCCC"):
    tcpr = cell._tc.get_or_add_tcPr(); tb = OxmlElement('w:tcBorders')
    for b in ['top','bottom','left','right']:
        e = OxmlElement(f'w:{b}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),'4'); e.set(qn('w:color'),color); tb.append(e)
    tcpr.append(tb)
def para(t, size=11, bold=False, italic=False, color=None, align=None, sb=4, sa=4):
    p = doc.add_paragraph(); p.paragraph_format.space_before=Pt(sb); p.paragraph_format.space_after=Pt(sa)
    if align: p.alignment=align
    r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(size)
    if bold: r.bold=True
    if italic: r.italic=True
    if color: r.font.color.rgb=color
    return p
def H(t, lvl=1):
    sz,co,sb,sa = (16,DARK,18,10) if lvl==1 else (13,DARK,14,8) if lvl==2 else (12,GREY,10,6)
    para(t, size=sz, bold=True, color=co, sb=sb, sa=sa)
def bullet(t, size=11):
    p=doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(size)
def num(t, size=11):
    p=doc.add_paragraph(style='List Number'); p.paragraph_format.space_before=Pt(2); p.paragraph_format.space_after=Pt(2)
    r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(size)
def formula(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(8); p.paragraph_format.space_after=Pt(8)
    r=p.add_run(t); r.font.name='Courier New'; r.font.size=Pt(12); r.bold=True; r.font.color.rgb=DARK
def caption(t):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_after=Pt(10)
    r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
def img(fn, w=6.5):
    p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.space_before=Pt(10); p.paragraph_format.space_after=Pt(5)
    fp=PROC/fn
    if fp.exists(): p.add_run().add_picture(str(fp), width=Inches(w))
def pb(): doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
def hdr(table, heads, widths=None):
    row=table.rows[0]
    for i,(c,t) in enumerate(zip(row.cells, heads)):
        if widths: c.width=Inches(widths[i])
        c.text=''; p=c.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(t); r.font.name='Arial'; r.font.size=Pt(10); r.bold=True; r.font.color.rgb=WHITE
        shade(c,HDR); borders(c); c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
def drow(table, vals, bold=None, colors=None, ac=True, bg=None):
    row=table.add_row(); bold=bold or [False]*len(vals); colors=colors or [None]*len(vals)
    for i,(c,t) in enumerate(zip(row.cells, vals)):
        c.text=''; p=c.paragraphs[0]
        if ac and i>0: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        r=p.add_run(str(t)); r.font.name='Arial'; r.font.size=Pt(10)
        if bold[i]: r.bold=True
        if colors[i]: r.font.color.rgb=colors[i]
        borders(c); c.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        if bg: shade(c,bg)

# ── TITTELSIDE ──
for _ in range(4): para("", 11)
para("Feltspesifikk Realisert Oljepris-modell", size=27, bold=True, color=DARK, align=WD_ALIGN_PARAGRAPH.CENTER, sb=20, sa=10)
para("Brent-linket differensial-dekomponering for norsk sokkel", size=17, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
para("Versjon 1.0 — med Script 63 kvalitets-fallback", size=12, italic=True, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER, sa=30)
for _ in range(3): para("", 11)
para("Emma Strandenskaar", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
para("Metodikk-dokumentasjon", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, sa=4)
para("Juni 2026", size=12, color=LGREY, align=WD_ALIGN_PARAGRAPH.CENTER, sa=30)
for _ in range(2): para("", 11)
t = doc.add_table(rows=1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Komponent","Metrikk","Verdi"], [2.6,2.4,1.4])
drow(t, ["Realisert pris (validering)","R² vs AKRBP rapportert","0.989"], bold=[True,False,True], colors=[None,None,GREEN])
drow(t, ["Realisert pris (validering)","MAE","1.56 USD/bbl"], colors=[None,None,GREEN])
drow(t, ["Differensial-modell","Out-of-time R²","0.379"])
drow(t, ["Differensial-modell","RMSE","2.78 USD/bbl"])
drow(t, ["Felt-dekning","Aker BP-felt med offisiell assay","12 / 12"])
drow(t, ["Felt-dekning","NCS-felt via Script 63 fallback","104"])
pb()

# ── 1. SAMMENDRAG ──
H("1. Sammendrag", 1)
para("Dette dokumentet beskriver en modell som dekomponerer Aker BPs (og generelt enhver NCS-operatørs) realiserte oljepris ned til feltnivå. Modellen forklarer hvorfor selskapets realiserte pris avviker fra Brent — drevet av oljekvalitet (API, svovel, metaller), logistikk og markedsforhold per felt.")
para("Kjerne-identiteten er enkel:", bold=True)
formula("Realisert_pris = Brent + Σᵢ [andelᵢ × differensialᵢ]")
para("der andelᵢ er feltets produksjonsandel og differensialᵢ er den modell-predikerte prisforskjellen mot Brent for feltets crude-kvalitet i gitt marked.")
para("Modellen validerer mot Aker BPs faktiske rapporterte realiserte priser med R² = 0.989 og MAE 1.56 USD/bbl over 24 kvartaler (2020-2026).")
para("Nytt i denne versjonen:", bold=True)
para("Felt uten offisiell publisert assay henter nå kvalitetsvektor automatisk fra Script 63 (NCS felt→kvalitet-imputering) i stedet for en ren geografi-fallback. Dette gjør modellen portabel til andre operatører og fremtidige feltutbygginger (f.eks. Yggdrasil-komponentene Frøy/Fulla/Munin).")

# ── 2. KONSEPT ──
H("2. Konseptet: realisert pris som Brent pluss differensial", 1)
para("Et oljeselskaps realiserte pris er ikke Brent. Hvert felt produserer en crude med spesifikk kvalitet som prises med et påslag eller en rabatt mot Brent-referansen:")
bullet("Lett, søt olje (høy API, lav svovel) → premium mot Brent (lettere å raffinere til høyverdiprodukter)")
bullet("Tung, sur olje (lav API, høy svovel) → rabatt mot Brent")
bullet("Logistikk: avstand til marked, FPSO vs rørledning → påvirker netto realisering")
para("Selskapets blended realiserte pris er produksjonsvektet snitt av feltenes differensialer, lagt til Brent. Når produksjonsmiksen endres (mix-shift), endres realisert pris selv om Brent er konstant — dette er kjernen i ER-analysen.")

# ── 3. TO-TRINNS STRUKTUR ──
H("3. To-trinns modellstruktur", 1)
H("3.1 Trinn 1 — Grade-nivå differensial-regresjon", 2)
para("En parsimonisk OLS-regresjon predikerer differensialen mot Brent for hver crude-grade, fra statiske kvalitetsvariabler interagert med dynamiske markedsforhold. Trent på 2 207 grade-måned observasjoner fra normpris-differensialer.")
para("Modellen kombinerer:")
bullet("Statisk kvalitet: API gravity, svovel, CCR, vacuum resid, vanadium/nikkel, middeldestillat-utbytte")
bullet("Logistikk: avstand-til-marked, FPSO-flagg, region (Nordsjøen)")
bullet("Dynamisk marked: Brent-nivå, diesel/bensin-cracks, Brent-Dubai-spread, raffineri-utnyttelse, lagernivåer, forward-kurve")
bullet("Interaksjoner: svovel×Brent, vacuum-resid×Brent, CCR×Brent (kvalitetsrabatten skalerer med oljeprisnivå)")
bullet("Politiske dummies: Russland-/Iran-sanksjoner, COVID, OPEC+-kutt")

H("3.2 Trinn 2 — Feltnivå produksjonsvektet blending", 2)
para("For hvert felt × måned bygges full feature-matrise (feltets kvalitet × det måneds markedet), differensialen predikeres, og kvartalsvis blandes feltene med Sodir-produksjonsvekter (justert for Aker BPs eierandel per felt).")

# ── 4. DIFFERENSIAL-MODELLEN ──
H("4. Differensial-modellen — drivere og presisjon", 1)
para("Modell: Parsimonisk OLS (Brent-linket, klyngerobuste standardfeil), 25 features.")
t = doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Metrikk","Verdi"], [3.3,3.3])
drow(t, ["In-sample R²","0.492"], ac=False)
drow(t, ["Cross-validated R²","0.472"], ac=False)
drow(t, ["Out-of-time R² (holdt ut nyeste periode)","0.379"], ac=False)
drow(t, ["RMSE","2.78 USD/bbl"], ac=False)
drow(t, ["Observasjoner","2 207 grade-måneder"], ac=False)
caption("Tabell: Differensial-modellens ytelse. OOT R²=0.38 er det ærlige out-of-time-anslaget.")

H("4.1 Hoveddrivere", 2)
para("De viktigste kvalitetsdriverne av differensialen (med markedsinteraksjoner):")
bullet("Svovel — sterkest kvalitetsdriver; rabatten skalerer med Brent-nivå (svovel×Brent-interaksjon) og raffineri-utnyttelse")
bullet("API gravity — lettere olje gir premium")
bullet("CCR / vacuum resid — bunnfraksjon trekker ned verdien (mer lavverdi-rest)")
bullet("Logistikk: lang avstand (−6.3 USD/bbl) og FPSO-lasting (−2.0 USD/bbl) gir rabatt")
bullet("Marked: raffineri-slakk, sanksjoner og OPEC+-kutt forskyver hele differensial-nivået")

# ── 5. DATAKILDER ──
H("5. Datakilder", 1)
t = doc.add_table(rows=1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Kilde","Bruk","Rolle"], [2.4,2.8,1.4])
drow(t, ["Normpris-differensialer","Treningsmål for differensial-regresjon","Target"], ac=False)
drow(t, ["Unified crude assays","Statisk kvalitet per grade (Equinor/Exxon/Total)","Feature"], ac=False)
drow(t, ["Sodir månedlig produksjon","Produksjonsvekter per felt × eierandel","Vekt"], ac=False)
drow(t, ["Markedspanel","Brent, cracks, spreads, lager, sanksjoner","Feature"], ac=False)
drow(t, ["AKRBP kvartalsrapporter","Validering mot rapportert realisert pris","Validering"], ac=False)
drow(t, ["Script 63 felt→kvalitet","Fallback for felt uten offisiell assay","Fallback"], ac=False)

# ── 6. FELT → KVALITET MED SCRIPT 63 FALLBACK ──
H("6. Felt → kvalitet med Script 63 fallback (nytt)", 1)
para("Hjertet av feltspesifisiteten er hvordan hvert felt tildeles en kvalitetsvektor. Modellen bruker nå et eksplisitt prioritert hierarki:")
t = doc.add_table(rows=1, cols=3); t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Prioritet","Kilde","Når"], [1.4,2.8,2.4])
drow(t, ["1 (høyest)","Hardkodet offisiell assay","Felt med publisert Equinor-assay eller navngitt eksport-blend"], ac=False, colors=[GREEN,None,None])
drow(t, ["2","Script 63: standalone","Feltet har egen publisert assay"], ac=False)
drow(t, ["3","Script 63: blend","Feltet selges som navngitt blend m/ assay"], ac=False)
drow(t, ["4","Script 63: median+DST","Ingen assay: område-median + feltspesifikk Sodir-DST API"], ac=False, colors=[ORANGE,None,None])

para("VIKTIG NYANSE — pris ≠ decline:", bold=True)
para("For PRIS teller salgsvaren. Et felt som selges inn i en navngitt blend (f.eks. Edvard Grieg → Grane Blend) prises som blenden, uansett reservoar-API. Derfor er blend-assayen korrekt prisinput — det motsatte av decline-modellen, der reservoar-API var riktig.")
para("Sodir wellbore (DST) gir kun API/GOR/temperatur, ikke svovel/metaller/destillasjonskutt — de krever lab-assay. Derfor: for felt uten assay hentes API fra Sodir-DST (feltspesifikk, målt), mens lab-variablene tas fra område-median (validert robust valg — API forutsier ikke svovel/metaller innen NCS).")

img("63_field_quality_imputation.png", 6.5)
caption("Figur 1: Script 63 felt→kvalitet-imputering. 104 NCS-felt med eksplisitt provenance-tier. Område-median for lab-variabler slår NCS-bred median (+25% svovel, +37% CCR); API fra Sodir-DST er den feltspesifikke gevinsten.")

H("6.1 Effekt på Aker BP-modellen", 2)
para("Alle 12 produserende Aker BP-felt har hardkodet offisiell assay (tier 1) — så den validerte modellen er uendret (R²=0.989). Fallbacken aktiverer for felt utenfor dette settet, og gjør modellen portabel:")
bullet("Yggdrasil-komponenter (Frøy, Fulla, Munin, Symra) resolver automatisk til tier-3 median+DST-vektorer — klar for forward-looking analyse")
bullet("Andre operatørers felt (Equinor-satellitter, Vår Energi) kan analyseres uten manuell assay-innlegging")

# ── 7. VALIDERING ──
H("7. Validering mot rapporterte realiserte priser", 1)
para("Modellen valideres mot Aker BPs faktiske rapporterte realiserte priser (kvartalsrapporter, 'realised price liquids'), hentet direkte fra selskapets investor relations.")
t = doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Valideringsmetrikk","Verdi"], [3.3,3.3])
drow(t, ["R² vs rapportert","0.989"], ac=False, colors=[None,GREEN])
drow(t, ["Korrelasjon","0.995"], ac=False)
drow(t, ["MAE","1.56 USD/bbl"], ac=False, colors=[None,GREEN])
drow(t, ["RMSE","2.02 USD/bbl"], ac=False)
drow(t, ["Bias","+0.30 USD/bbl"], ac=False)
drow(t, ["Antall kvartaler","24 (2020-2026)"], ac=False)
caption("Tabell: Validering. Lav bias (+0.30) bekrefter at modellen fanger AKRBP realisert pris strukturelt, ikke bare i nivå.")

img("42_akrbp_realized_decomposition.png", 6.5)
caption("Figur 2: Aker BP realisert pris-dekomponering. Øverst: modell-predikert vs rapportert realisert pris. Midten: feltbidrag til differensial. Nederst: produksjonsandeler og kvalitetsposisjonering per felt.")

# ── 8. BRUK I ER ──
H("8. Bruk i Equity Research", 1)
bullet("Dekomponering: hvilke felt driver realisert pris opp/ned, og med hvor mye")
bullet("Mix-shift-prognose: når produksjonsmiksen endres (Yggdrasil kommer inn 2027), endres realisert pris — kvantifiserbart")
bullet("Følsomhet: hvordan realisert pris responderer på Brent-nivå, cracks, sanksjoner per felt")
bullet("Cross-selskap: samme rammeverk på Equinor, Vår Energi via Script 63-fallback")
para("Differansen mot konsensus: de fleste analytikere bruker en flat Brent-rabatt for hele selskapet. Denne modellen gir feltspesifikk dekomponering med en validert R²=0.989 mot faktiske rapporterte tall.")

# ── 9. BEGRENSNINGER ──
H("9. Begrensninger", 1)
bullet("Differensial-modellens OOT R²=0.38 — markedsdynamikken har betydelig støy; realisert-pris-valideringen (R²=0.989) er høyere fordi produksjonsvekting og Brent-nivå dominerer")
bullet("Sodir-produksjonslag: nyeste kvartal kan mangle data (mars-spike ikke fanget) — ekskluderes fra metrikk")
bullet("NGL/kondensat: rapportert 'realised liquids' inkluderer NGL; modellen approksimerer dette via grade-miks")
bullet("Tier-3-felt (median+DST): lab-variabler har reell restusikkerhet (svovel ±0.06%-poeng); brukes kun for felt uten assay, og har lav produksjonsvekt for Aker BP")
bullet("Blend-tilordning forutsetter kjent eksportrute; nye felt kan endre rute over tid")

# ── APPENDIX ──
pb()
H("Appendix A: Differensial-modell hovedkoeffisienter", 1)
t = doc.add_table(rows=1, cols=2); t.alignment=WD_TABLE_ALIGNMENT.CENTER
hdr(t, ["Variabel","Koeffisient (USD/bbl)"], [4.0,2.6])
for k,v in [("Konstant","−18.58"),("Svovel %","+13.58"),("Lang avstand (dummy)","−6.28"),
            ("Region Nordsjøen","−3.56"),("CCR vekt-%","−3.19"),("FPSO-lasting","−1.98"),
            ("Raffineri-slakk","+1.29"),("COVID","−1.27"),("OPEC+-kutt 2023","−1.10"),
            ("Russland-sanksjoner","+1.02"),("API gravity","+0.87"),("Svovel×Brent","−0.085")]:
    drow(t, [k,v], ac=False)
caption("Koeffisientene tolkes i sammenheng med interaksjoner; svovel-effekten skaleres f.eks. ned med svovel×Brent ved høye oljepriser.")

H("Appendix B: Variabel-definisjoner", 1)
para("Differensial: prisforskjell (USD/bbl) mellom en crude-grade og Brent-referansen.")
para("Produksjonsandel: feltets netto oljeproduksjon (× eierandel) / total netto produksjon i kvartalet.")
para("Realisert pris: produksjonsvektet salgspris selskapet faktisk oppnår, inkl. NGL.")
para("Tier (Script 63): provenance for feltets kvalitetsvektor — standalone > blend > median+DST.")
para("OOT R²: out-of-time — modellen trent på eldre periode, testet på holdt-ut nyeste periode.")

H("Appendix C: Prosjektmateriale", 1)
para("Hovedscript: scripts/42_akrbp_realized_price_decomposition.py")
para("Differensial-modell: data/processed/34b_brent_model.json")
para("Felt→kvalitet fallback: scripts/63_field_quality_imputation.py → 63_ncs_field_quality.csv")
para("Alt versjonskontrollert i prosjektets git-repository.")

OUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUT))
print(f"Lagret: {OUT}  ({OUT.stat().st_size/1024:.0f} KB)")
try:
    shutil.copy2(OUT, ONEDRIVE); print(f"Kopiert til OneDrive: {ONEDRIVE}")
except Exception as e:
    print(f"OneDrive-kopi feilet: {e}")
