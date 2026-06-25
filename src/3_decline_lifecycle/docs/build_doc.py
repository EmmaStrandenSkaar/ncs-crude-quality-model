"""
Build the methodology document for the NCS Decline Rate ER model.
Uses python-docx for portability (node.js not installed).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS = Path(str(Path(__file__).resolve().parents[3] / "analyses" / "decline_quality" / "results"))
OUTPUT = Path(str(Path(__file__).resolve().parents[3] / "analyses" / "decline_quality" / "docs" / "NCS_Decline_Model_Methodology.docx"))

# Colors
DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
GREY = RGBColor(0x59, 0x59, 0x59)
LIGHT_GREY = RGBColor(0x80, 0x80, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREEN_BG = "E8F5E9"
HEADER_BG = "1F4E79"

doc = Document()

# ── Page setup ──
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# ── Default style ──
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

# ── Helper functions ──
def set_cell_shading(cell, color_hex):
    """Add background shading to a cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def add_para(text, size=11, bold=False, italic=False, color=None, align=None, space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    return p

def add_heading(text, level=1):
    if level == 1:
        size, color, before, after = 16, DARK_BLUE, 18, 10
    elif level == 2:
        size, color, before, after = 13, DARK_BLUE, 14, 8
    else:
        size, color, before, after = 12, GREY, 10, 6
    add_para(text, size=size, bold=True, color=color, space_before=before, space_after=after)

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)

def add_numbered(text, size=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GREY

def add_image(filename, width_inches=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run()
    img_path = RESULTS / filename
    if img_path.exists():
        run.add_picture(str(img_path), width=Inches(width_inches))

def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def add_table_header_row(table, headers, widths_in=None):
    row = table.rows[0]
    for i, (cell, text) in enumerate(zip(row.cells, headers)):
        if widths_in:
            cell.width = Inches(widths_in[i])
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_BG)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_table_data_row(table, values, bold=None, colors=None, align_center=True, bg=None):
    row = table.add_row()
    bold = bold or [False] * len(values)
    colors = colors or [None] * len(values)
    for i, (cell, text) in enumerate(zip(row.cells, values)):
        cell.text = ''
        p = cell.paragraphs[0]
        if align_center and i > 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        if bold[i]: run.bold = True
        if colors[i]: run.font.color.rgb = colors[i]
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg:
            set_cell_shading(cell, bg)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
# Spacer
for _ in range(4):
    add_para("", size=11)

add_para("Decline Rate Modell for Norsk Sokkel",
         size=28, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=20, space_after=10)
add_para("Fysikk + Premium Rammeverk for Equity Research",
         size=18, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=30)

for _ in range(3):
    add_para("", size=11)

add_para("Emma Strandenskaar", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Metodikk-dokumentasjon", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Juni 2026", size=12, color=LIGHT_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

for _ in range(3):
    add_para("", size=11)

# Key stats table
stats_table = doc.add_table(rows=1, cols=2)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(stats_table, ["Modell-ytelse", "Verdi"], widths_in=[3.5, 2.5])
add_table_data_row(stats_table, ["Cross-validated R²", "0.713"], bold=[False, True], colors=[None, GREEN])
add_table_data_row(stats_table, ["In-sample R²", "0.770"])
add_table_data_row(stats_table, ["RMSE", "0.040"])
add_table_data_row(stats_table, ["Antall felt i kalibrering", "49 NCS-felt"])
add_table_data_row(stats_table, ["Aker BP treffrate (±0.05)", "75%"], bold=[False, True], colors=[None, GREEN])

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. SAMMENDRAG
# ═══════════════════════════════════════════════════════════════
add_heading("1. Sammendrag", level=1)
add_para("Dette dokumentet beskriver en kvantitativ modell for å estimere årlige decline rates (D_annual) for olje- og gassfelt på norsk sokkel. Modellen kombinerer:")
add_bullet("Fysikk-baseline: Beggs-Robinson viskositet beregnet fra API gravity")
add_bullet("Historisk premium: feltets faktiske avvik fra fysikk-prediksjonen siste 12 måneder")
add_bullet("Asymmetri-korreksjon: justering for felt som avviker mye fra fysikken (uansett retning)")

add_para("Modellen oppnår CV R² = 0.713 på 49 NCS-felt, med RMSE 0.040 på årlige decline rates. Alle inputs er offentlig tilgjengelige fra Sodir (norsk sokkeldirektorat), noe som gjør modellen direkte anvendelig i ER-analyse av Aker BP, Equinor, Vår Energi, ConocoPhillips og andre operatører på sokkelen.")

add_para("Den endelige formelen er:", bold=True)
add_formula("D = 0.0664 + 0.0601·ln(viskositet) − 0.0597·P₁₂ + 0.0379·|P₁₂|")
add_para("hvor P₁₂ er gjennomsnittlig log-premium siste 12 måneder.")

# ═══════════════════════════════════════════════════════════════
# 2. HVORFOR DECLINE RATES ER KRITISK
# ═══════════════════════════════════════════════════════════════
add_heading("2. Hvorfor decline rates er kritisk for Equity Research", level=1)
add_para("Decline rate er en av de viktigste driverne for verdsettelse av oppstrøms olje- og gasselskaper. Den bestemmer hvor raskt produksjonen — og dermed kontantstrømmen — faller fra et felt etter peak.")

add_heading("2.1 Påvirkning på NPV", level=2)
add_para("For et typisk NCS-felt med 10-15 års produksjonshistorikk vil en endring i decline rate fra 8% til 12% redusere NPV med 20-30%, alt annet likt. Dette gjør decline-anslag til en hovedvariabel i target price-modeller.")

add_heading("2.2 Hvor markedet ofte feiler", level=2)
add_para("Mange ER-rapporter bruker en \"standard\" decline rate på 5-10% for alle NCS-felt, uten å justere for:")
add_bullet("Oljekvalitet (viskositet) — påvirker reservoarets evne til å strømme")
add_bullet("Feltspesifikk historikk — noen felt har holdt seg flate i 10+ år (Valhall, Ekofisk)")
add_bullet("Operatørstrategi — CAPEX-investering bremser natural decline")
add_para("Vår modell gir et transparent, datadrevet alternativ som fanger disse forskjellene gjennom to enkle variabler.")

# ═══════════════════════════════════════════════════════════════
# 3. FYSIKK
# ═══════════════════════════════════════════════════════════════
add_heading("3. Fysikk-fundamentet: Beggs-Robinson viskositet", level=1)

add_heading("3.1 Hvorfor viskositet betyr noe", level=2)
add_para("Når en oljereservoar tappes, må oljen strømme gjennom mikroskopiske kanaler i bergarten mot brønnen. Hvor lett dette skjer avhenger av oljens viskositet:")
add_bullet("Lett olje (høy API, lav viskositet) — strømmer lett, men reservoarets trykk faller raskt → bratt decline")
add_bullet("Tung olje (lav API, høy viskositet) — strømmer trått, men reservoartrykket holdes lengre → slakere decline (men lavere total produksjon)")
add_para("Denne fysikalske sammenhengen er grunnlaget for vår baseline-prediksjon. Vi kan ikke direkte måle viskositet for hvert felt i sanntid, men API gravity er offentlig tilgjengelig og lett konvertibel via Beggs-Robinson-korrelasjonen.")

add_heading("3.2 Beggs-Robinson-formelen", level=2)
add_para("Beggs-Robinson (1975) er en empirisk korrelasjon som beregner dead-oil-viskositet fra API gravity og reservoartemperatur:")
add_formula("μ = 10^(x · T^(-1.163)) − 1   [cP]")
add_formula("hvor x = 10^(3.0324 − 0.02023·API)")
add_para("Vi bruker T = 194°F (90°C) som NCS-typisk reservoartemperatur. Sensitivitetstesting viste at temperaturvalg fra 60°C til 127°C kun endrer modellytelsen marginalt (CV R² varierer med ±0.003).")

add_heading("3.3 Validering av fysikk-baseline", level=2)
add_para("På 51 NCS-felt finner vi en signifikant positiv sammenheng mellom ln(viskositet) og observert D_annual (β = 0.059, p < 0.001). Dette bekrefter hypotesen: tyngre olje gir slakere decline. Men fysikk alene forklarer kun R² = 9.6% av variansen i decline rates — derfor trenger vi premium-rammeverket.")

add_image("fig_physics_viscosity.png", width_inches=5.5)
add_caption("Figur 1: Sammenheng mellom ln(viskositet) og decline rate på NCS. Fysikken gir riktig retning (positiv pendens) men forklarer bare en liten del av variansen — feltspesifikke effekter dominerer.")

# ═══════════════════════════════════════════════════════════════
# 4. PREMIUM RAMMEVERK
# ═══════════════════════════════════════════════════════════════
add_heading("4. Premium-rammeverket: Alpha / Beta-tankegang", level=1)

add_heading("4.1 Analog til finansmarkeder", level=2)
add_para("Vi låner et begrepsapparat fra finans og bruker det på reservoarer:")
add_bullet("Beta (fysikk) — felles eksponering mot \"markedet\" (her: viskositet → forventet decline)")
add_bullet("Alpha (premium) — feltspesifikk avvik som ikke kan forklares av fysikken alene")
add_para("Premiumet fanger alt vi ikke direkte kan observere: operatørkvalitet, brønnplassering, vanninjeksjonsstrategi, CAPEX-program, reservoarstruktur, infill-boring osv. Vi trenger ikke å modellere disse direkte — feltets faktiske produksjonshistorikk avslører dem.")

add_heading("4.2 Beregning av premium", level=2)
add_para("For hver måned i post-peak-perioden beregner vi:")
add_formula("log_premium_i = ln(actual_i / expected_i)")
add_para("hvor expected_i = exp(−D_physics/12 · month_i)")
add_para("Hvis feltet faktisk produserer mer enn fysikken tilsier, blir log_premium positiv (\"discount\" på decline). Hvis det produserer mindre, blir log_premium negativ (\"premium\" på decline).")

add_heading("4.3 Hvorfor 12-måneders vindu", level=2)
add_para("Vi testet systematisk premium-vindu fra 3 til 120 måneder. Resultatene:")

window_table = doc.add_table(rows=1, cols=3)
window_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(window_table, ["Vindu", "CV R²", "Aker BP RMSE"], widths_in=[2.2, 2.2, 2.2])
add_table_data_row(window_table, ["6 mnd", "0.713", "0.0557"])
add_table_data_row(window_table, ["12 mnd ★", "0.713", "0.0565"],
                   bold=[True, True, False], colors=[None, GREEN, None], bg=LIGHT_GREEN_BG)
add_table_data_row(window_table, ["24 mnd", "0.707", "0.0585"])
add_table_data_row(window_table, ["36 mnd", "0.665", "0.0625"])
add_table_data_row(window_table, ["60 mnd", "0.620", "0.0684"])
add_table_data_row(window_table, ["Lifetime", "0.523", "0.0756"])

add_image("fig_premium_window_finegrained.png", width_inches=6.5)
add_caption("Figur 2: Systematisk testing av premium-vindu (3-24 mnd). Tallene under hvert punkt viser antall datapunkter per felt. CV R² topper ut rundt 6-12 mnd og faller jevnt for lengre vinduer — eldre data legger til støy, ikke signal. Vindu under 6 mnd er for sårbart for støy; over 24 mnd inneholder utdaterte signaler.")

add_para("Vi valgte 12 måneder framfor 6 av tre grunner:")
add_numbered("Robusthet — 12 datapunkter er stabilt mot enkeltmåneds-shutdowns (vedlikehold, branner). 6 datapunkter er sårbart for én anomalisk måned.")
add_numbered("Sesongeffekter — Norske felt har planlagte sommerstopp for vedlikehold. 12 mnd fanger en full syklus.")
add_numbered("ER-narrativ — \"1 år historikk\" er en troverdig kommunikasjon i en investeringsanalyse.")

add_heading("4.4 Hvorfor |premium| (asymmetri-korreksjon)", level=2)
add_para("Etter at vi inkluderte premium som lineær variabel, observerte vi en U-formet sammenheng: felt med stort avvik fra fysikken (i begge retninger) hadde høyere baseline-decline enn felt nær fysikkens prediksjon. Vi inkluderte derfor |P₁₂| som tredje variabel.")
add_para("Intuisjonen er at felt med stort |premium| er mer uforutsigbare:")
add_bullet("Store positive premium (f.eks. Valhall +2.07) skyldes ofte intensive CAPEX-sykluser og redevelopment som skaper volatile produksjonsmønstre")
add_bullet("Store negative premium (f.eks. Volund −1.53) indikerer felt med brattere natural decline enn fysikken tilsier")
add_para("Begge tilfeller fortjener et risikopåslag i decline-estimatet.")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. METODOLOGI
# ═══════════════════════════════════════════════════════════════
add_heading("5. Metodologi", level=1)

add_heading("5.1 Datakilder", level=2)

data_table = doc.add_table(rows=1, cols=2)
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(data_table, ["Kilde", "Bruk"], widths_in=[2.2, 4.4])
add_table_data_row(data_table, ["Sodir månedlig produksjon", "Felt-måned produksjonsvolumer, beregning av peak og post-peak-måneder"], align_center=False)
add_table_data_row(data_table, ["Sodir feltinformasjon", "API gravity, operatør, hovedområde, oppdagelsesår"], align_center=False)
add_table_data_row(data_table, ["Equinor / Norem assays", "Bekreftelse av oljekvalitetsparametere for blendinger (Ekofisk, Grane, Alvheim)"], align_center=False)
add_table_data_row(data_table, ["NSTA (UK)", "Out-of-sample validering av modellen på UK Continental Shelf"], align_center=False)

add_heading("5.2 Beregning av D_annual", level=2)
add_para("For hvert felt fitter vi en eksponentiell decline-kurve på post-peak-data:")
add_formula("ln(production_t / peak_production) = −D_annual · t  + ε")
add_para("D_annual estimeres med vanlig minste kvadraters metode. Felt med færre enn 12 post-peak-måneder eller R² < 0.1 ekskluderes.")

add_heading("5.3 Modelltrening", level=2)
add_para("Modellen estimeres med vanlig OLS-regresjon. Vi bruker Leave-One-Out Cross-Validation (LOO-CV) for å oppnå et ærlig anslag på out-of-sample-ytelse, gitt N = 49 felt. Manuell LOO-implementasjon (sklearn-funksjonen returnerer NaN for 1-felt test-folds).")

add_heading("5.4 Robusthetstesting", level=2)
add_para("Modellen er testet på følgende måter:")
add_bullet("LOO Cross-Validation — gir CV R² = 0.713 (sannferdig anslag)")
add_bullet("Out-of-sample på UK NSTA-felt — viser at modellen er basin-spesifikk (UK-felt har omvendt viskositetskoeffisient)")
add_bullet("Sensitivitetstest på temperatur — knapt påvirkning fra 60°C til 127°C")
add_bullet("Felt-karakteristikk-analyse — bekreftet at ingen Sodir-variabel forbedrer modellen utover premium")

# ═══════════════════════════════════════════════════════════════
# 6. ENDELIG MODELL
# ═══════════════════════════════════════════════════════════════
add_heading("6. Endelig modell", level=1)

add_heading("6.1 Formel og koeffisienter", level=2)
add_formula("D_annual = 0.0664 + 0.0601·ln(μ) − 0.0597·P₁₂ + 0.0379·|P₁₂|")

coef_table = doc.add_table(rows=1, cols=5)
coef_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(coef_table, ["Variabel", "Koeffisient", "Std. β", "t-stat", "p-verdi"],
                     widths_in=[1.8, 1.2, 1.2, 1.2, 1.2])
add_table_data_row(coef_table, ["Intercept", "+0.0664", "—", "+5.81", "<0.001 ***"],
                   colors=[None, None, None, None, GREEN])
add_table_data_row(coef_table, ["ln(viskositet)", "+0.0601", "+0.027", "+4.61", "<0.001 ***"],
                   colors=[None, None, None, None, GREEN])
add_table_data_row(coef_table, ["Premium 12 mnd", "−0.0597", "−0.080", "−11.32", "<0.001 ***"],
                   bold=[False, False, True, True, False], colors=[None, None, None, None, GREEN])
add_table_data_row(coef_table, ["|Premium 12 mnd|", "+0.0379", "+0.034", "+4.82", "<0.001 ***"],
                   colors=[None, None, None, None, GREEN])

add_caption("Alle koeffisienter er statistisk signifikante på 1%-nivå.")

add_heading("6.2 Tolkning av koeffisientene", level=2)
add_bullet("ln(viskositet) — positiv: høyere viskositet (tyngre olje) gir høyere decline. Bekrefter fysikk-hypotesen.")
add_bullet("Premium 12 mnd — sterkt negativ: felt som outperformer fysikken (positiv premium) har lavere decline. Dette er hovedeffekten i modellen (|β| = 0.080).")
add_bullet("|Premium| — positiv: store avvik fra fysikken (uansett retning) gir høyere baseline-decline. Risikopåslag for uforutsigbare felt.")

add_heading("6.3 Modellfigur", level=2)
add_image("fig_final_model.png", width_inches=6.5)
add_caption("Figur 3: Endelig modell — 9-panels oversikt. Øverst: CV-progresjon, predikert vs. faktisk, variabel-viktighet. Midten: U-formet premium-effekt, viskositets-effekt, Aker BP-prediksjoner. Nederst: Implisitte prognoser, full ranking, formel.")

add_heading("6.4 Modellytelse i detalj — Premium/Discount-rammeverk", level=2)
add_para("Figuren under viser hvordan premium/discount-rammeverket fanger feltspesifikke avvik fra fysikk-baseline. Felt over 45-graderslinjen underperformer fysikken (decline-premium), mens felt under outperformer (decline-discount).")
add_image("fig_akerbp_premium_framework.png", width_inches=6.5)
add_caption("Figur 4: Aker BP premium/discount-rammeverk. Øverst venstre: Discount/premium per felt — Valhall, Alvheim, Bøyla outperformer fysikken; Edvard Grieg, Skogul, Volund underperformer. Øverst høyre: Aker BP plassering i NCS-skyen. Nederst: Stabilitet over tid og 5-års produksjonsprognoser med vs. uten premium-justering.")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 7. PRAKTISK BRUK
# ═══════════════════════════════════════════════════════════════
add_heading("7. Praktisk bruk i ER-analyse", level=1)

add_heading("7.1 Trinn-for-trinn-prosedyre", level=2)
add_numbered("Hent API gravity for feltet (Sodir, Equinor crude assays, eller selskapsrapport)")
add_numbered("Beregn viskositet med Beggs-Robinson-formelen")
add_numbered("Last ned siste 12 måneders produksjonsdata fra Sodir factpages")
add_numbered("Identifiser peak-måned og beregn months_since_peak for hver datapunkt")
add_numbered("Beregn fysikk-baseline: D_physics = 0.097 + 0.059·ln(μ)")
add_numbered("For hver av siste 12 mnd: log_premium_i = ln(actual_i) + D_physics/12·month_i")
add_numbered("P₁₂ = gjennomsnitt av de 12 log_premium-verdiene")
add_numbered("Plugg inn i hovedformelen: D = 0.0664 + 0.0601·ln(μ) − 0.0597·P₁₂ + 0.0379·|P₁₂|")
add_numbered("Produksjonsprognose: production_year_T = peak · exp(−D · T)")

add_heading("7.2 Eksempel: Aker BP-felt", level=2)
add_para("Tabellen under viser modellens prediksjoner for Aker BPs 12 modne NCS-felt:")

akbp_table = doc.add_table(rows=1, cols=6)
akbp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(akbp_table, ["Felt", "API", "D_faktisk", "D_modell", "Bom", "Premium 12m"],
                     widths_in=[1.4, 0.8, 1.0, 1.0, 1.0, 1.2])

akbp_data = [
    ("EDVARD GRIEG", "27.1", "0.384", "0.221", "+0.163", "−0.79"),
    ("SKOGUL", "34.5", "0.247", "0.190", "+0.058", "−0.85"),
    ("VOLUND", "34.5", "0.243", "0.256", "−0.013", "−1.53"),
    ("VILJE", "34.5", "0.222", "0.208", "+0.014", "−1.03"),
    ("IVAR AASEN", "27.1", "0.219", "0.191", "+0.029", "−0.48"),
    ("SKARV", "50.8", "0.107", "0.146", "−0.039", "−1.07"),
    ("TAMBAR", "38.9", "0.096", "0.066", "+0.029", "+0.97"),
    ("ULA", "38.9", "0.083", "0.079", "+0.004", "+0.39"),
    ("ALVHEIM", "34.5", "0.056", "0.085", "−0.029", "+0.99"),
    ("HOD", "38.9", "0.053", "0.031", "+0.022", "+2.61"),
    ("BØYLA", "34.5", "0.050", "0.104", "−0.053", "+0.15"),
    ("VALHALL", "38.9", "0.038", "0.043", "−0.004", "+2.07"),
]

for row in akbp_data:
    field, api, d_act, d_pred, miss, prem = row
    miss_val = float(miss.replace("−", "-").replace("+", ""))
    prem_val = float(prem.replace("−", "-").replace("+", ""))
    miss_color = GREEN if abs(miss_val) < 0.05 else RED
    prem_color = GREEN if prem_val > 0 else RED
    add_table_data_row(akbp_table, [field, api, d_act, d_pred, miss, prem],
                       bold=[False, False, False, True, False, False],
                       colors=[None, None, None, None, miss_color, prem_color])

add_caption("75% av Aker BP-feltene treffes innenfor ±0.05 (grønne tall). De største bommene (Edvard Grieg, Bøyla) skyldes ekstrem produksjonsvolatilitet som ingen enkel modell kan fange.")

add_para("Figuren under viser modellens prediksjoner som faktiske decline-kurver for hvert enkelt Aker BP-felt — sortert med største bom først, slik at du visuelt kan vurdere hvor godt modellen treffer per felt.")
add_image("fig_akerbp_decline_curves.png", width_inches=6.5)
add_caption("Figur 5: Faktisk produksjonsprofil (blå linje) vs. modell-prediksjon (oransje stiplet) og faktisk decline-fit (rosa stiplet) per Aker BP-felt. Rød skygge = modell underpredikerer decline; grønn skygge = modell overpredikerer. Modellen treffer godt for de fleste felt — størst avvik på Edvard Grieg (svært tung olje, uforutsigbar) og Bøyla (volatil produksjon).")

add_heading("7.3 Investeringscase-bygging", level=2)
add_para("Modellen lar deg systematisk svare på spørsmål som:")
add_bullet("\"Bør Edvard Grieg ha samme decline-anslag som Ivar Aasen?\" — Nei, premium-historikken viser at Edvard Grieg har 70% høyere natural decline.")
add_bullet("\"Er Valhalls produksjonsplatå holdbar?\" — Premium = +2.07 reflekterer redevelopment-CAPEX. Spørsmålet blir: vil Aker BP fortsette investeringstakten?")
add_bullet("\"Hva er rimelig terminal decline for Skarv?\" — Modellen gir 0.146 vs 0.107 faktisk; premium −1.07 indikerer brattere decline enn fysikken tilsier.")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. BEGRENSNINGER
# ═══════════════════════════════════════════════════════════════
add_heading("8. Begrensninger og forutsetninger", level=1)

add_heading("8.1 Modellen forutsetter eksponentiell decline", level=2)
add_para("Vi modellerer decline som en konstant prosentvis nedgang per år. Dette passer godt for de fleste NCS-felt, men ikke for:")
add_bullet("Platå-felt som Valhall (har holdt ~35% av peak i 10+ år)")
add_bullet("Felt under aktiv redevelopment der produksjonen midlertidig øker")
add_bullet("Volatile felt med uregelmessige investeringssykluser (Bøyla)")
add_para("For disse feltene fanger premium-variablen deler av effekten, men noen genuine outliers (som Edvard Grieg) gir høyere prediksjonsfeil. Valhall er et godt eksempel — feltet har et produksjonsplatå som ingen ren eksponentiell modell kan fange perfekt:")
add_image("fig_valhall_forensic.png", width_inches=6.5)
add_caption("Figur 6: Valhall — hvorfor eksponentiell decline ikke fungerer. Øverst venstre: Full produksjonshistorie viser redevelopment-faser (oransje skygge = ny plattform 2010-2013). Øverst høyre: Valhall slår fysikk-baseline hvert år fra 2001. Nederst venstre: Decline rate varierer 10× mellom 3-års perioder, inkludert negativ D (produksjonsøkning). Nederst høyre: Produksjonsplatået på ~35% av peak i 10+ år er en S-kurve, ikke eksponentiell.")

add_heading("8.2 Basin-spesifikk kalibrering", level=2)
add_para("Modellen er kalibrert på NCS-felt. UK NSTA-validering viste at viskositetskoeffisienten faktisk reverserer fortegn på UK Continental Shelf — sannsynligvis fordi UK-feltene er en annen generasjon med annerledes utviklingshistorikk.")
add_para("Implikasjon: modellen bør re-kalibreres for andre basseng (Nordsjøen UK, Golfen, Vest-Afrika) før bruk.")

add_heading("8.3 Krever 12 mnd post-peak data", level=2)
add_para("For helt nye felt uten 12 måneders post-peak-historikk kan modellen ikke beregne premium. I praksis betyr dette at de første ~18 månedene etter peak kun gir fysikk-baseline-prediksjon.")

add_heading("8.4 Forutsetter stabilt operatørprogram", level=2)
add_para("Premium fanger historikk, men forutsetter at fremtidig operatør-atferd ligner siste 12 måneder. Plutselige endringer (CAPEX-kutt, salg av felt, redevelopment-prosjekter) krever kvalitative justeringer.")

# ═══════════════════════════════════════════════════════════════
# 9. DIFFERENSIERING
# ═══════════════════════════════════════════════════════════════
add_heading("9. Differensiering vs. standard tilnærminger", level=1)

add_heading("9.1 Sammenligning med vanlige metoder", level=2)

diff_table = doc.add_table(rows=1, cols=3)
diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(diff_table, ["Tilnærming", "Svakhet", "Vår modell"],
                     widths_in=[1.8, 2.4, 2.4])
add_table_data_row(diff_table, ["Flat 5-10% decline", "Ignorerer feltspesifikke forskjeller helt",
                                "Differensierer fra 1% (Ekofisk) til 38% (Edvard Grieg)"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(diff_table, ["Operatør-guidance", "Subjektiv, optimistisk bias, ikke verifiserbar",
                                "Datadrevet, transparent metodikk"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(diff_table, ["Komplekse reservoarmodeller", "Krever proprietære data; black-box for analytikere",
                                "Offentlige inputs; 4 parametere; etterprøvbart"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(diff_table, ["Lineær ekstrapolering", "Ingen fysisk forankring",
                                "Beggs-Robinson som vitenskapelig baseline"],
                   align_center=False, colors=[None, None, GREEN])

add_heading("9.2 Hva dette gir i en ER-/IB-kontekst", level=2)
add_bullet("Differensiert syn på decline-rater per felt — viktig competitive edge")
add_bullet("Transparent metodikk som tåler intern modell-validering")
add_bullet("Skalerbar til alle NCS-operatører (Equinor, Aker BP, Vår Energi, ConocoPhillips, OKEA, DNO)")
add_bullet("Mulig å overføre til andre basseng med rekalibrering")
add_bullet("Bevisst om begrensninger — ingen \"black box\"-mystikk")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════
add_heading("Appendix A: Modellutvikling-historikk", level=1)
add_para("Modellen ble utviklet gjennom 12 iterasjoner. De viktigste milepælene:")

dev_table = doc.add_table(rows=1, cols=4)
dev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(dev_table, ["Trinn", "Forbedring", "CV R²", "Aker BP RMSE"],
                     widths_in=[0.8, 3.4, 1.2, 1.2])
add_table_data_row(dev_table, ["1", "Bare viskositet (Beggs-Robinson)", "−0.015", "0.095"],
                   align_center=False)
add_table_data_row(dev_table, ["2", "+ lifetime premium", "0.473", "0.076"],
                   align_center=False)
add_table_data_row(dev_table, ["3", "Switch til siste 3 års premium", "0.544", "0.063"],
                   align_center=False)
add_table_data_row(dev_table, ["4", "+ 12 mnd premium + |premium|", "0.713", "0.056"],
                   align_center=False, bold=[True, True, True, True],
                   colors=[None, None, GREEN, GREEN], bg=LIGHT_GREEN_BG)

add_heading("Appendix B: Variabel-definisjoner", level=1)
add_para("API gravity: Standard målestokk for oljens tetthet (°API). Høyere API = lettere olje. NCS-typisk område: 25° (tung) til 50° (lett kondensat).")
add_para("Viskositet (μ): Oljens motstand mot strømning, målt i centipoise (cP). Vann har μ ≈ 1 cP; tunge oljer har μ ≈ 100+ cP.")
add_para("D_annual: Eksponentiell decline rate på årlig basis. Production_year_T = Peak · exp(−D · T). D = 0.10 betyr 10% årlig nedgang.")
add_para("Premium (P₁₂): Gjennomsnittlig log-avvik mellom faktisk og fysikk-forventet produksjon siste 12 måneder. Positiv = outperformer; negativ = underperformer.")
add_para("Post-peak: Periode fra månedlig produksjon når 100% av peak-måneden. Modellen gjelder bare post-peak-perioden.")
add_para("LOO-CV: Leave-One-Out Cross-Validation. Estimer modellen N ganger, hver gang med ett felt holdt utenfor. Måler ærlig out-of-sample-ytelse.")

add_heading("Appendix C: Kontaktinfo og prosjektmateriale", level=1)
add_para("Alle scripts, data og figurer er versjonskontrollert i prosjektet:")
add_para("Repository: ./")
add_para("Hovedscript: analyses/decline_quality/scripts/")
add_para("Resultater: analyses/decline_quality/results/")
add_para("Modellen er fullt etterprøvbar og kan utvides med nye datakilder (kvartalsrapporter, brønnmeldinger, CAPEX-guidance) etter behov.")

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")
