"""
Build the methodology document v2 with reservoir-level fluid enrichment + V5 model.
"""

from pathlib import Path
import shutil
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RESULTS = Path(str(Path(__file__).resolve().parents[3] / "analyses" / "decline_quality" / "results"))
DOCS = Path(str(Path(__file__).resolve().parents[3] / "analyses" / "decline_quality" / "docs"))
OUTPUT = DOCS / "NCS_Decline_Model_Methodology.docx"  # overwrite v2
ONEDRIVE = Path(str(Path.home()) + "/Library/CloudStorage/OneDrive-BINorwegianBusinessSchool(BIEDU)/NCS_Decline_Model_Methodology.docx")

DARK_BLUE = RGBColor(0x1F, 0x4E, 0x79)
GREEN = RGBColor(0x2E, 0x7D, 0x32)
RED = RGBColor(0xC6, 0x28, 0x28)
ORANGE = RGBColor(0xE6, 0x51, 0x00)
GREY = RGBColor(0x59, 0x59, 0x59)
LIGHT_GREY = RGBColor(0x80, 0x80, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_GREEN_BG = "E8F5E9"
LIGHT_BLUE_BG = "E3F2FD"
LIGHT_ORANGE_BG = "FFF3E0"
HEADER_BG = "1F4E79"

doc = Document()
section = doc.sections[0]
section.page_height = Inches(11)
section.page_width = Inches(8.5)
section.left_margin = Inches(1)
section.right_margin = Inches(1)
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(11)

def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for border_name in ['top', 'bottom', 'left', 'right']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:color'), color)
        tc_borders.append(border)
    tc_pr.append(tc_borders)

def add_para(text, size=11, bold=False, italic=False, color=None, align=None, space_before=4, space_after=4):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align: p.alignment = align
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    return p

def add_heading(text, level=1):
    if level == 1: size, color, before, after = 16, DARK_BLUE, 18, 10
    elif level == 2: size, color, before, after = 13, DARK_BLUE, 14, 8
    else: size, color, before, after = 12, GREY, 10, 6
    add_para(text, size=size, bold=True, color=color, space_before=before, space_after=after)

def add_bullet(text, size=11):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)

def add_numbered(text, size=11):
    p = doc.add_paragraph(style='List Number')
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(size)

def add_formula(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(12)
    run.bold = True
    run.font.color.rgb = DARK_BLUE

def add_caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(text)
    run.font.name = 'Arial'
    run.font.size = Pt(9)
    run.italic = True
    run.font.color.rgb = GREY

def add_image(filename, width_inches=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(5)
    run = p.add_run()
    img_path = RESULTS / filename
    if img_path.exists():
        run.add_picture(str(img_path), width=Inches(width_inches))

def add_page_break():
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)

def add_table_header_row(table, headers, widths_in=None):
    row = table.rows[0]
    for i, (cell, text) in enumerate(zip(row.cells, headers)):
        if widths_in: cell.width = Inches(widths_in[i])
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        run.bold = True
        run.font.color.rgb = WHITE
        set_cell_shading(cell, HEADER_BG)
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

def add_table_data_row(table, values, bold=None, colors=None, align_center=True, bg=None):
    row = table.add_row()
    bold = bold or [False] * len(values)
    colors = colors or [None] * len(values)
    for i, (cell, text) in enumerate(zip(row.cells, values)):
        cell.text = ''
        p = cell.paragraphs[0]
        if align_center and i > 0: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(str(text))
        run.font.name = 'Arial'
        run.font.size = Pt(10)
        if bold[i]: run.bold = True
        if colors[i]: run.font.color.rgb = colors[i]
        set_cell_borders(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if bg: set_cell_shading(cell, bg)

# ═══════════════════════════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════════════════════════
for _ in range(4):
    add_para("", size=11)

add_para("Decline Rate Modell for Norsk Sokkel",
         size=28, bold=True, color=DARK_BLUE, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_before=20, space_after=10)
add_para("Fysikk + Premium Rammeverk for Equity Research",
         size=18, italic=True, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=4)
add_para("Versjon 3.0 — full lifecycle (decline + pre-peak forecast)",
         size=12, italic=True, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=30)

for _ in range(3):
    add_para("", size=11)

add_para("Emma Strandenskaar", size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Metodikk-dokumentasjon", size=12, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
add_para("Juni 2026 — Revidert utgave", size=12, color=LIGHT_GREY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=30)

for _ in range(2):
    add_para("", size=11)

stats_table = doc.add_table(rows=1, cols=3)
stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(stats_table, ["Motor", "Nøkkelmetrikk", "Verdi"], widths_in=[2.4, 3.0, 2.0])
add_table_data_row(stats_table, ["Decline (V5.1)", "Nested CV R²", "0.662"], align_center=False, bold=[True, False, True], colors=[None, None, GREEN])
add_table_data_row(stats_table, ["Decline (V5.1)", "Aker BP treffrate (±0.05)", "83% [55-95% CI]"], align_center=False)
add_table_data_row(stats_table, ["Pre-peak forecast", "Ramp/platå (out-of-sample)", "±5/±4 mnd"], align_center=False, bold=[True, False, True], colors=[None, None, GREEN])
add_table_data_row(stats_table, ["Pre-peak forecast", "Peak (out-of-sample)", "±35% (range)"], align_center=False, colors=[None, None, ORANGE])
add_table_data_row(stats_table, ["Datagrunnlag", "Master fluid library", "108 felt"], align_center=False)
add_table_data_row(stats_table, ["Datagrunnlag", "Type-curve bibliotek", "69 felt"], align_center=False)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 1. SAMMENDRAG
# ═══════════════════════════════════════════════════════════════
add_heading("1. Sammendrag", level=1)
add_para("Dette dokumentet beskriver et komplett produksjons-forecasting-rammeverk for norsk sokkel, bygget rundt to komplementære motorer:")
add_bullet("Decline-modell (V5.1) — for felt med ≥12 måneders produksjonshistorikk. Estimerer årlig decline rate fra fysikk (Beggs-Robinson viskositet) + et feltspesifikt premium-rammeverk. Seksjon 2-10.")
add_bullet("Pre-peak forecast (V2) — for helt nye felt, fra PDO-data alene. Predikerer peak / ramp / platå / decline fra kun ex-ante variabler. Seksjon 11-13.")

add_para("Sammen dekker de hele feltets livssyklus: fra førsteolje (pre-peak forecast) til moden produksjon (decline-modell). For et nytt felt brukes pre-peak forecast til feltet har 12 måneders historikk, deretter tar den mer presise decline-modellen over.")

add_para("DECLINE-MODELLEN (V5.1):", bold=True)
add_para("Bruker reservoar-API fra Sodir DST og operatør-research i stedet for handelsblending-API. Oppnår ærlig nested CV R² = 0.662 og Aker BP treffrate 83% (innen ±0.05). Alle inputs offentlig tilgjengelige.")
add_formula("D = 0.0938 + 0.0106·ln(viskositet) − 0.0612·P₁₂ + 0.0399·|P₁₂|")

add_para("PRE-PEAK FORECAST (V2):", bold=True)
add_para("Predikerer ramp/platå presist (±5/±4 mnd out-of-sample), peak som range-verktøy (±35%). Validert med ekte hold-out test der modellen aldri så testfeltet. Brukt på Yggdrasil som case study (seksjon 13).")


# ═══════════════════════════════════════════════════════════════
# 1.5 QA-FIKSER (NY SEKSJON)
# ═══════════════════════════════════════════════════════════════
add_heading("1.5 QA-fikser anvendt i V5.1", level=1)
add_para("V5.1 inkorporerer 7 fikser fra en grundig kvalitetsrevisjon:")

qa_table = doc.add_table(rows=1, cols=3)
qa_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(qa_table, ["Fix", "Problem", "Løsning"], widths_in=[0.7, 3.0, 2.9])
add_table_data_row(qa_table, ["H1", "TROLL API = 38.8° fra operator (kondensat-blend)",
                              "Erstattet med Sodir DST oljerim 27.9° (n=15)"], align_center=False)
add_table_data_row(qa_table, ["H2", "Duplikater: ASGARD/ÅSGARD, KVITEBJORN/KVITEBJØRN",
                              "Norsk Ø/O normalisering for dedup → 108 unike felt"], align_center=False)
add_table_data_row(qa_table, ["H3", "Gass-/kondensatfelt forurenser oljemodell",
                              "DUVA, SNØHVIT, ORMEN LANGE flagget som gas_field"], align_center=False)
add_table_data_row(qa_table, ["H5", "Enkelt-DST sample overstyrer blend",
                              "Krev n≥3 før DST aksepteres; ellers nedgrades til dst_single_unreliable"], align_center=False)
add_table_data_row(qa_table, ["M1", "Premium-vindu off-by-one: 13 mnd i stedet for 12",
                              "Endret t≥t.max()-12 til t>t.max()-12 (korrekte 12 obs)"], align_center=False)
add_table_data_row(qa_table, ["A1/A2", "Premium-sirkularitet: bruker D_physics som er fittet på D_annual",
                              "Implementert nested LOO-CV: refit både D_physics og premium per fold"], align_center=False)
add_table_data_row(qa_table, ["Stat", "Bootstrap CIs + Wilson hit-rate CI",
                              "Rapporterer nå 95% CIs for alle koeffisienter og proporsjoner"], align_center=False)

add_para("Effekt på rapportert ytelse:", bold=True)

effect_table = doc.add_table(rows=1, cols=3)
effect_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(effect_table, ["Metrikk", "V5 (gammel)", "V5.1 (ærlig)"], widths_in=[2.5, 2.0, 2.0])
add_table_data_row(effect_table, ["CV R² (LOO)", "0.702 (inflert)", "0.662 (nested)"],
                   colors=[None, RED, GREEN])
add_table_data_row(effect_table, ["Premium-vindu", "13 mnd (bug)", "12 mnd (korrekt)"])
add_table_data_row(effect_table, ["Aker BP hit-rate", "83%", "83% [55-95% CI]"])
add_table_data_row(effect_table, ["N felt (etter dedup)", "49", "48"])

add_caption("Tabell: Største endring er CV R² 0.702 → 0.662. Selv om dette er en nedgang, er det det ÆRLIGE out-of-sample-anslaget. Modellen er fortsatt sterkt positiv vs fysikk-only baseline (~0).")

add_para("Hvorfor denne nedgangen er bra:", bold=True)
add_bullet("Den gamle CV R² = 0.702 var inflert pga sirkulær konstruksjon av premium-variabelen")
add_bullet("V5.1 rapporterer det ekte out-of-sample-tallet (0.662) — modellen er fortsatt nyttig")
add_bullet("Premium-rammeverket bidrar 65pp R² over fysikk-only baseline — fortsatt sterkt signal")
add_bullet("ER-rapporten kan nå quote 0.662 trygt uten å overselge")

add_para("hvor viskositet beregnes med Beggs-Robinson på reservoar-API (når tilgjengelig) og felt-spesifikk reservoartemperatur (der målt), og P₁₂ er gjennomsnittlig log-premium siste 12 måneder.")

# ═══════════════════════════════════════════════════════════════
# 2. HVORFOR DECLINE RATES ER KRITISK
# ═══════════════════════════════════════════════════════════════
add_heading("2. Hvorfor decline rates er kritisk for Equity Research", level=1)
add_para("Decline rate er en av de viktigste driverne for verdsettelse av oppstrøms olje- og gasselskaper. Den bestemmer hvor raskt produksjonen — og dermed kontantstrømmen — faller fra et felt etter peak.")

add_heading("2.1 Påvirkning på NPV", level=2)
add_para("For et typisk NCS-felt med 10-15 års produksjonshistorikk vil en endring i decline rate fra 8% til 12% redusere NPV med 20-30%, alt annet likt. Dette gjør decline-anslag til en hovedvariabel i target price-modeller.")

add_heading("2.2 Hvor markedet ofte feiler", level=2)
add_para("Mange ER-rapporter bruker en \"standard\" decline rate på 5-10% for alle NCS-felt, uten å justere for:")
add_bullet("Oljekvalitet (reservoar-viskositet) — påvirker reservoarets evne til å strømme")
add_bullet("Feltspesifikk historikk — noen felt har holdt seg flate i 10+ år (Valhall, Ekofisk)")
add_bullet("Operatørstrategi — CAPEX-investering bremser natural decline")
add_para("Vår modell gir et transparent, datadrevet alternativ som fanger disse forskjellene gjennom få variabler.")

# ═══════════════════════════════════════════════════════════════
# 3. FYSIKK
# ═══════════════════════════════════════════════════════════════
add_heading("3. Fysikk-fundamentet: Beggs-Robinson viskositet", level=1)

add_heading("3.1 Hvorfor viskositet betyr noe", level=2)
add_para("Når en oljereservoar tappes, må oljen strømme gjennom mikroskopiske kanaler i bergarten mot brønnen. Hvor lett dette skjer avhenger av oljens viskositet:")
add_bullet("Lett olje (høy API, lav viskositet) — strømmer lett, men reservoarets trykk faller raskt → bratt decline")
add_bullet("Tung olje (lav API, høy viskositet) — strømmer trått, men reservoartrykket holdes lengre → slakere decline")

add_heading("3.2 Beggs-Robinson-formelen", level=2)
add_para("Beggs-Robinson (1975) er en empirisk korrelasjon som beregner dead-oil-viskositet fra API gravity og reservoartemperatur:")
add_formula("μ = 10^(x · T^(-1.163)) − 1   [cP]")
add_formula("hvor x = 10^(3.0324 − 0.02023·API)")
add_para("I V2 av modellen bruker vi felt-spesifikk reservoartemperatur for 10 felt der DST-målinger eller operatør-rapporter har dette. For øvrige felt brukes NCS-typisk T = 90°C (194°F). Sensitivitetstesting viste at temperaturvalg har marginal effekt på modell-ytelse — men felt-spesifikk T er mer ærlig modellering.")

add_heading("3.3 Validering av fysikk-baseline", level=2)
add_image("fig_physics_viscosity.png", width_inches=5.5)
add_caption("Figur 1: Sammenheng mellom ln(viskositet) og decline rate på NCS. Fysikken gir riktig retning men forklarer alene bare en liten del av variansen — feltspesifikke effekter dominerer.")

# ═══════════════════════════════════════════════════════════════
# 4. PREMIUM-RAMMEVERKET
# ═══════════════════════════════════════════════════════════════
add_heading("4. Premium-rammeverket: Alpha / Beta-tankegang", level=1)

add_heading("4.1 Analog til finansmarkeder", level=2)
add_para("Vi låner et begrepsapparat fra finans og bruker det på reservoarer:")
add_bullet("Beta (fysikk) — felles eksponering mot \"markedet\" (her: viskositet → forventet decline)")
add_bullet("Alpha (premium) — feltspesifikk avvik som ikke kan forklares av fysikken alene")
add_para("Premiumet fanger alt vi ikke direkte kan observere: operatørkvalitet, brønnplassering, vanninjeksjonsstrategi, CAPEX-program, reservoarstruktur, infill-boring osv.")

add_heading("4.2 Beregning av premium", level=2)
add_para("For hver måned i post-peak-perioden beregner vi:")
add_formula("log_premium_i = ln(actual_i / expected_i)")
add_para("hvor expected_i = exp(−D_physics/12 · month_i)")
add_para("Positiv premium = feltet outperformer fysikken. Negativ premium = feltet underperformer.")

add_heading("4.3 Hvorfor 12-måneders vindu", level=2)
add_para("Vi testet systematisk premium-vindu fra 3 til 120 måneder. CV R² topper ut rundt 6-12 måneder.")
add_image("fig_premium_window_finegrained.png", width_inches=6.5)
add_caption("Figur 2: Systematisk testing av premium-vindu (3-24 mnd). Tallene under hvert punkt viser antall datapunkter per felt. Lengre vinduer legger til støy, ikke signal. Vindu under 6 mnd er for sårbart for støy.")

add_para("Vi valgte 12 måneder framfor 6 av tre grunner:")
add_numbered("Robusthet — 12 datapunkter er stabilt mot enkeltmåneds-shutdowns.")
add_numbered("Sesongeffekter — Norske felt har planlagte sommerstopp som 12 mnd fanger en full syklus av.")
add_numbered("ER-narrativ — \"1 år historikk\" er en troverdig kommunikasjon i en investeringsanalyse.")

add_heading("4.4 Hvorfor |premium| (asymmetri-korreksjon)", level=2)
add_para("Felt med stort avvik fra fysikken (i begge retninger) har høyere baseline-decline enn felt nær fysikkens prediksjon. Vi inkluderer derfor |P₁₂| som tredje variabel.")
add_bullet("Store positive premium (f.eks. Valhall +2.34) skyldes intensive CAPEX-sykluser")
add_bullet("Store negative premium (f.eks. Volund −1.64) indikerer brattere natural decline")
add_para("Begge tilfeller fortjener et risikopåslag i decline-estimatet.")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 5. RESERVOAR-API BERIKELSE (NY SEKSJON!)
# ═══════════════════════════════════════════════════════════════
add_heading("5. Reservoar-API berikelse (nytt i v2)", level=1)

add_heading("5.1 Problemet: blend-API vs reservoar-API", level=2)
add_para("I første versjon av modellen brukte vi API gravity fra handelsblendinger (f.eks. \"Alvheim Blend\", \"Grane Blend\"). Dette er API-en på oljen som selges fra hub-systemet — ikke API-en på den faktiske reservoarvæsken.")

add_para("Problemet er at satellittfelt rør oljen til hovedfeltets FPSO/plattform der den blandes med hovedfeltets olje. Det som selges er blendingen, og API-målingen gjøres på blendingen. Sodir tildeler samme blend-API til alle felt som matter inn samme strøm:", italic=True)

# Blend examples table
blend_table = doc.add_table(rows=1, cols=3)
blend_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(blend_table, ["Blending", "Felt med samme API-tildelt", "Blend-API"],
                     widths_in=[2.0, 4.0, 1.4])
add_table_data_row(blend_table, ["Alvheim", "ALVHEIM, VILJE, VOLUND, BØYLA, SKOGUL", "34.5°"], align_center=False)
add_table_data_row(blend_table, ["Ekofisk", "EKOFISK, HOD, ELDFISK, EMBLA, ULA, VALHALL, GYDA, TAMBAR", "38.9°"], align_center=False)
add_table_data_row(blend_table, ["Grane", "GRANE, EDVARD GRIEG, IVAR AASEN, SOLVEIG", "27.1°"], align_center=False)
add_table_data_row(blend_table, ["Statfjord", "STATFJORD, SNORRE, SYGNA, STATFJORD ØST/NORD, TORDIS, VIGDIS", "39.3°"], align_center=False)
add_caption("Tabell: Eksempler på blend-tildelinger som ble brukt i v1. Disse satellittfeltene har faktisk forskjellige reservoarvæsker.")

add_heading("5.2 Løsningen: Sodir DST-database + operatør-research", level=2)
add_para("For å bygge en database med reservoar-API per felt, kombinerte vi flere datakilder:")

source_table = doc.add_table(rows=1, cols=3)
source_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(source_table, ["Kilde", "Type data", "Felt dekket"],
                     widths_in=[2.5, 3.0, 1.5])
add_table_data_row(source_table, ["Sodir DST-database", "1186 brønntester med oljedensitet (g/cm³) per oppdagelsesbrønn", "85"], align_center=False)
add_table_data_row(source_table, ["Operatør direct assays", "Equinor, ExxonMobil, TotalEnergies — komplette crude assays", "17"], align_center=False)
add_table_data_row(source_table, ["Operatør deep research", "Annual reports, capital markets days, PDOer", "43"], align_center=False)
add_table_data_row(source_table, ["Sodir field_summary (v1)", "Eksisterende blend-tildelinger", "52"], align_center=False)

add_caption("Tabell: Datakilder for master fluid library. Total unik dekning: 110 felt.")

add_heading("5.3 Prioritert merge-logikk", level=2)
add_para("For hvert felt velger vi den beste tilgjengelige API-verdien etter denne prioritetsrekkefølgen:")
add_numbered("Operatør direct assay (høy konfidens) — kvalitet: operator_direct")
add_numbered("Sodir DST med ≥5 brønntester (median av samples) — kvalitet: dst_robust")
add_numbered("Operatør medium-konfidens research (med formasjon, depth, etc.) — kvalitet: operator_medium")
add_numbered("Sodir DST med 1-4 brønntester — kvalitet: dst_limited")
add_numbered("Operatør lav-konfidens research — kvalitet: operator_low")
add_numbered("Blend-tildeling som fallback — kvalitet: blend_inherited")

add_heading("5.4 Største endringer fra v1", level=2)
add_para("Av 49 modellerte felt har 31 fått oppdatert API. De største endringene:")

change_table = doc.add_table(rows=1, cols=4)
change_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(change_table, ["Felt", "v1 API (blend)", "v2 API (reservoar)", "Kilde"],
                     widths_in=[2.0, 1.5, 1.5, 2.0])
for row in [
    ("EDVARD GRIEG", "27.1°", "36.5°", "Aker BP research"),
    ("KNARR", "26.0°", "47.9°", "Sodir DST"),
    ("MARULK", "34.0°", "48.6°", "Sodir DST"),
    ("MARIA", "24.4°", "47.7°", "Sodir DST"),
    ("MARTIN LINGE", "27.8°", "43.5°", "Sodir DST"),
    ("SKARV", "50.8°", "43.3°", "Aker BP direct"),
    ("ÅSGARD", "47.5°", "41.1°", "Sodir DST (n=47)"),
    ("TROLL", "38.8°", "27.9°", "Sodir DST (n=15)"),
    ("OSEBERG", "39.9°", "36.1°", "Sodir DST (n=34)"),
    ("GULLFAKS", "39.3°", "31.7°", "Sodir DST (n=23)"),
    ("EKOFISK", "38.9°", "35.1°", "Sodir DST (n=20)"),
    ("BALDER", "30.0°", "25.6°", "Sodir DST (n=10)"),
]:
    add_table_data_row(change_table, list(row),
                       colors=[None, None, GREEN if "DST" in row[3] or "direct" in row[3] else None, None])
add_caption("Tabell: Største API-endringer fra blend til reservoar. Avg. absolutt endring på tvers av 31 felt: 3.9°.")

add_heading("5.5 Den ærlige avveiingen: v3 vs v5", level=2)
add_para("Vi testet 5 modell-versjoner og fant en interessant avveiing mellom presisjon og fysikalsk ærlighet:")

versions_table = doc.add_table(rows=1, cols=5)
versions_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(versions_table, ["Versjon", "Beskrivelse", "CV R²", "Aker BP RMSE", "ln(μ) β"],
                     widths_in=[1.0, 2.8, 1.0, 1.4, 1.2])
add_table_data_row(versions_table, ["V1", "Original (blend API)", "0.713", "0.056", "+0.060"])
add_table_data_row(versions_table, ["V2", "Full reservoir-API replacement", "0.701", "0.064", "−0.053"])
add_table_data_row(versions_table, ["V3", "Hybrid: kun høy-konfidens", "0.713", "0.056", "+0.053"])
add_table_data_row(versions_table, ["V4", "Hybrid: høy + medium", "0.702", "0.062", "+0.018"])
add_table_data_row(versions_table, ["V5 ★", "Hybrid + felt-spesifikk T", "0.702", "0.061", "+0.011"],
                   bold=[True, True, True, True, True],
                   colors=[GREEN, GREEN, GREEN, GREEN, GREEN], bg=LIGHT_GREEN_BG)
add_caption("Tabell: 5 modell-versjoner. V3 har best CV R², men V5 valgt for fysikalsk ærlighet.")

add_image("fig_hybrid_comparison.png", width_inches=6.5)
add_caption("Figur 3: Sammenligning av 5 modell-versjoner. V3 vinner på CV R², men V5 valgt som anbefalt modell pga. ærligere reservoarmodellering. Nederste panel viser at de fleste Aker BP felt har like små residualer på tvers av versjoner.")

add_para("Vi valgte V5 som anbefalt modell selv om V1/V3 har marginalt høyere CV R². Grunnene er:", bold=True)
add_bullet("Ærligere reservoarmodellering: bruker reservoarmålinger der vi har dem, ikke fiktive blend-tildelinger")
add_bullet("Bedre Aker BP treffrate: 83% innen ±0.05 (V5) vs 75% (V1)")
add_bullet("Mer generaliserbar: når vi predikterer nye felt har V5 et mer korrekt fysisk grunnlag")
add_bullet("Pitch-bart: \"Vi bruker reservoardata, ikke handelsdata\" er en sterk differensieringsfaktor")
add_bullet("Premium-rammeverket absorberer fortsatt feltspesifikke effekter, så vi mister minimal presisjon")

add_para("Dette er en bevisst metodologisk avveiing. V1/V3 var mer presise på vårt 49-felt sett, men V5 vil sannsynligvis være mer pålitelig out-of-sample.", italic=True)

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 6. METODOLOGI
# ═══════════════════════════════════════════════════════════════
add_heading("6. Metodologi", level=1)

add_heading("6.1 Datakilder", level=2)
data_table = doc.add_table(rows=1, cols=2)
data_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(data_table, ["Kilde", "Bruk"], widths_in=[2.2, 4.4])
add_table_data_row(data_table, ["Sodir månedlig produksjon", "Felt-måned produksjonsvolumer, beregning av peak og post-peak-måneder"], align_center=False)
add_table_data_row(data_table, ["Sodir DST-database", "Reservoar-API fra brønntester (oljedensitet g/cm³)"], align_center=False)
add_table_data_row(data_table, ["Sodir wellbore exploration", "Mapping av brønner til felt for DST-aggregering"], align_center=False)
add_table_data_row(data_table, ["Equinor crude assays", "Direct assay for 17 felt (Statfjord, Gullfaks, Oseberg, Troll, etc.)"], align_center=False)
add_table_data_row(data_table, ["Aker BP capital markets day", "Reservoar-detaljer for 17 felt"], align_center=False)
add_table_data_row(data_table, ["Vår Energi / OKEA / Harbour / DNO", "Operatør-specific reservoar-data"], align_center=False)

add_heading("6.2 Beregning av D_annual", level=2)
add_para("For hvert felt fitter vi en eksponentiell decline-kurve på post-peak-data:")
add_formula("ln(production_t / peak_production) = −D_annual · t + ε")

add_heading("6.3 Modelltrening", level=2)
add_para("Modellen estimeres med vanlig OLS-regresjon. Vi bruker Leave-One-Out Cross-Validation (LOO-CV) for å oppnå et ærlig anslag på out-of-sample-ytelse, gitt N = 49 felt.")

# ═══════════════════════════════════════════════════════════════
# 7. ENDELIG MODELL V5
# ═══════════════════════════════════════════════════════════════
add_heading("7. Endelig modell V5.1 (med QA-fikser)", level=1)

add_heading("7.1 Formel og koeffisienter", level=2)
add_formula("D_annual = 0.0938 + 0.0106·ln(μ) − 0.0612·P₁₂ + 0.0399·|P₁₂|")

coef_table = doc.add_table(rows=1, cols=5)
coef_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(coef_table, ["Variabel", "Koeffisient", "Std. β", "t-stat", "p-verdi"],
                     widths_in=[1.8, 1.2, 1.2, 1.2, 1.2])
add_table_data_row(coef_table, ["Intercept", "+0.0938", "—", "[+0.07, +0.11]", "<0.001 ***"],
                   colors=[None, None, None, None, GREEN])
add_table_data_row(coef_table, ["ln(viskositet)", "+0.0106", "+0.006", "[−0.01, +0.03]", "ns"],
                   colors=[None, None, None, None, GREY])
add_table_data_row(coef_table, ["Premium 12 mnd", "−0.0612", "−0.086", "[−0.08, −0.05]", "<0.001 ***"],
                   bold=[False, False, True, True, False], colors=[None, None, None, None, GREEN])
add_table_data_row(coef_table, ["|Premium 12 mnd|", "+0.0399", "+0.039", "[+0.03, +0.07]", "<0.001 ***"],
                   colors=[None, None, None, None, GREEN])
add_caption("Tabell: V5 koeffisienter. Premium-variablene er svært signifikante. ln(viskositet) er ikke signifikant alene — den fungerer som baseline-justering, og premium fanger feltspesifikke effekter.")

add_heading("7.2 Tolkning av koeffisientene", level=2)
add_bullet("ln(viskositet) — svakt positiv, ikke signifikant alene. Premium absorberer det meste av viskositets-informasjonen via feltets historiske avvik.")
add_bullet("Premium 12 mnd — sterkt negativ: felt som outperformer fysikken (positiv premium) har lavere decline. Dette er hovedeffekten (|β|=0.086).")
add_bullet("|Premium| — positiv: store avvik fra fysikken (uansett retning) gir høyere baseline-decline. Risikopåslag for uforutsigbare felt.")

add_heading("7.3 Modellfigur", level=2)
add_image("fig_final_v5_model.png", width_inches=6.5)
add_caption("Figur 4: V5 endelig modell. Øverst: predikert vs faktisk, residual-fordeling (sentrert), variabel-viktighet. Midten: API-oppdatering, U-formet premium-effekt, Aker BP-prediksjoner. Nederst: API-kvalitet per felt, datakilder, formel.")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 8. PRAKTISK BRUK
# ═══════════════════════════════════════════════════════════════
add_heading("8. Praktisk bruk i ER-analyse", level=1)

add_heading("8.1 Trinn-for-trinn-prosedyre", level=2)
add_numbered("Hent reservoar-API for feltet (sjekk Sodir DST, operatør CMD, eller assay rapporter — ikke handelsblend-API)")
add_numbered("Hent reservoartemperatur hvis tilgjengelig (ellers bruk 90°C som NCS-default)")
add_numbered("Beregn viskositet med Beggs-Robinson")
add_numbered("Last ned siste 12 måneders produksjonsdata fra Sodir")
add_numbered("Beregn fysikk-baseline: D_physics = 0.097 + 0.059·ln(μ) (kalibrert på 49 NCS-felt)")
add_numbered("For hver av siste 12 mnd: log_premium_i = ln(actual_i) + D_physics/12·month_i")
add_numbered("P₁₂ = gjennomsnitt av de 12 log_premium-verdiene")
add_numbered("Plugg inn i V5-formelen: D = 0.0938 + 0.0106·ln(μ) − 0.0612·P₁₂ + 0.0399·|P₁₂|")

add_heading("8.2 Aker BP-felt: V5 prediksjoner", level=2)

akbp_table = doc.add_table(rows=1, cols=7)
akbp_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(akbp_table, ["Felt", "API v5", "T brukt", "D_faktisk", "D_modell", "Bom", "Premium 12m"],
                     widths_in=[1.3, 0.8, 0.9, 0.9, 0.9, 0.9, 1.0])

akbp_data = [
    ("EDVARD GRIEG", "36.5", "90°C", "0.384", "0.200", "+0.184", "−0.99"),
    ("SKOGUL", "34.5", "90°C", "0.247", "0.191", "+0.057", "−0.88"),
    ("VOLUND", "34.5", "90°C", "0.243", "0.267", "−0.024", "−1.64"),
    ("VILJE", "34.5", "90°C", "0.222", "0.217", "+0.005", "−1.14"),
    ("IVAR AASEN", "27.1", "90°C", "0.219", "0.177", "+0.042", "−0.67"),
    ("SKARV", "43.3", "90°C", "0.107", "0.147", "−0.040", "−0.52"),
    ("TAMBAR", "40.0", "90°C", "0.096", "0.074", "+0.022", "+1.12"),
    ("ULA", "40.0", "90°C", "0.083", "0.085", "−0.001", "+0.60"),
    ("ALVHEIM", "33.9", "78°C", "0.056", "0.085", "−0.029", "+0.93"),
    ("HOD", "36.0", "90°C", "0.053", "0.038", "+0.015", "+2.94"),
    ("BØYLA", "34.5", "90°C", "0.050", "0.100", "−0.050", "+0.07"),
    ("VALHALL", "36.0", "90°C", "0.038", "0.051", "−0.013", "+2.34"),
]

for row in akbp_data:
    field, api, T, d_act, d_pred, miss, prem = row
    miss_val = float(miss.replace("−", "-").replace("+", ""))
    prem_val = float(prem.replace("−", "-").replace("+", ""))
    miss_color = GREEN if abs(miss_val) < 0.05 else ORANGE if abs(miss_val) < 0.1 else RED
    prem_color = GREEN if prem_val > 0 else RED
    T_bold = T != "90°C"
    add_table_data_row(akbp_table, list(row),
                       bold=[False, False, T_bold, False, True, False, False],
                       colors=[None, None, ORANGE if T_bold else None, None, None, miss_color, prem_color])

add_caption("Tabell: V5 prediksjoner for Aker BP. 83% innen ±0.05. Alvheim bruker felt-spesifikk T=78°C. Edvard Grieg fortsatt outlier — uforutsigbar produksjon kan ikke fanges av noen enkel modell.")

add_heading("8.3 Investeringscase-bygging", level=2)
add_para("Modellen lar deg systematisk svare på spørsmål som:")
add_bullet("\"Bør Edvard Grieg ha samme decline-anslag som Ivar Aasen?\" — Nei, premium-historikken viser at Edvard Grieg har 75% høyere natural decline.")
add_bullet("\"Er Valhalls produksjonsplatå holdbar?\" — Premium = +2.34 reflekterer redevelopment-CAPEX. Spørsmålet blir: vil Aker BP fortsette investeringstakten?")
add_bullet("\"Hva er rimelig terminal decline for Skarv?\" — Modellen gir 0.147 vs 0.107 faktisk; premium −0.52 indikerer moderat underperformance vs fysikk.")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 9. BEGRENSNINGER
# ═══════════════════════════════════════════════════════════════
add_heading("9. Begrensninger og forutsetninger", level=1)

add_heading("9.1 Modellen forutsetter eksponentiell decline", level=2)
add_para("Vi modellerer decline som en konstant prosentvis nedgang per år. Dette passer godt for de fleste NCS-felt, men ikke for:")
add_bullet("Platå-felt som Valhall (har holdt ~35% av peak i 10+ år)")
add_bullet("Felt under aktiv redevelopment der produksjonen midlertidig øker")
add_bullet("Volatile felt med uregelmessige investeringssykluser (Bøyla)")

add_image("fig_valhall_forensic.png", width_inches=6.5)
add_caption("Figur 5: Valhall — hvorfor eksponentiell decline ikke fungerer for platå-felt. Produksjonen har holdt seg på ~35% av peak i 10+ år pga redevelopment-CAPEX.")

add_heading("9.2 V5: ærlighet over marginal presisjon", level=2)
add_para("V5 har marginalt lavere CV R² (0.702 vs V1's 0.713). Vi har akseptert denne avveiingen for fire fordeler:")
add_numbered("Modellen er fysisk korrekt — ingen kunstige sammenhenger fra blend-tildelinger")
add_numbered("Modellen generaliserer bedre — out-of-sample (nye felt) får sannsynligvis mer pålitelig prediksjon")
add_numbered("Aker BP treffrate (±0.05) gikk OPP fra 75% til 83%")
add_numbered("Differensieringen i pitchen er sterkere — \"reservoardata\" høres bedre enn \"handelsdata\"")

add_heading("9.3 Basin-spesifikk kalibrering", level=2)
add_para("Modellen er kalibrert på NCS-felt. UK NSTA-validering viste at viskositetskoeffisienten reverserer fortegn på UK Continental Shelf. Modellen bør re-kalibreres for andre basseng.")

add_heading("9.4 Krever 12 mnd post-peak data", level=2)
add_para("For helt nye felt uten 12 måneders post-peak-historikk kan modellen ikke beregne premium. I praksis betyr dette at de første ~18 månedene etter peak kun gir fysikk-baseline-prediksjon.")

# ═══════════════════════════════════════════════════════════════
# 10. DIFFERENSIERING
# ═══════════════════════════════════════════════════════════════
add_heading("10. Differensiering vs. standard tilnærminger", level=1)

diff_table = doc.add_table(rows=1, cols=3)
diff_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(diff_table, ["Tilnærming", "Svakhet", "Vår V5-modell"],
                     widths_in=[2.0, 2.5, 2.5])
add_table_data_row(diff_table, ["Flat 5-10% decline", "Ignorerer feltspesifikke forskjeller helt",
                                "Differensierer fra 1% (Ekofisk) til 38% (Edvard Grieg)"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(diff_table, ["Operatør-guidance", "Subjektiv, optimistisk bias",
                                "Datadrevet, transparent metodikk"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(diff_table, ["Blend-API som proxy", "Mister reservoar-fysikk",
                                "Reservoar-API + felt-T fra Sodir DST og operatør-research"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(diff_table, ["Komplekse reservoarmodeller", "Krever proprietære data",
                                "Offentlige inputs; 4 parametere; etterprøvbart"],
                   align_center=False, colors=[None, None, GREEN])

add_heading("10.1 Hva V2 av modellen gir i en ER-/IB-kontekst", level=2)
add_bullet("110-felts fluid library med reservoar-API per felt — direkte anvendelig for ER")
add_bullet("Ærligere fysikk-modellering basert på reservoardata, ikke handelsdata")
add_bullet("Field-spesifikk reservoartemperatur der målt — mer presis Beggs-Robinson")
add_bullet("Aker BP treffrate 83% — modellen er pitch-klar")
add_bullet("Transparent metodikk dokumentert i 110+ sider med kode")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# 11. PRE-PEAK FORECAST (NY SEKSJON — V3)
# ═══════════════════════════════════════════════════════════════
add_heading("11. Pre-peak forecasting — produksjon for nye felt", level=1)

add_para("Decline-modellen (seksjon 1-10) krever minst 12 måneders produksjonshistorikk for å beregne premium. For helt nye felt — som ennå ikke produserer — trenger vi et eget rammeverk som forecaster fra data tilgjengelig FØR produksjon (PDO og discovery DST).")

add_para("Dette er en fundamentalt annen modelleringsoppgave. For modne felt er decline det eneste ukjente. For nye felt har vi fire ukjente som alle påvirker NPV:")
add_bullet("Peak-rate — hvor høyt produksjonen topper")
add_bullet("Ramp-up tid — fra førsteolje til peak")
add_bullet("Platå-lengde — hvor lenge peak holder seg")
add_bullet("Decline-rate — fallet etter platå")

add_heading("11.1 Kun ex-ante variabler", level=2)
add_para("Et avgjørende metodisk prinsipp: forecast-modellen bruker KUN variabler kjent før feltet starter produksjon. Det er fristende å inkludere variabler som forbedrer historisk fit, men hvis de ikke er kjent på forhånd er de ubrukelige for en ekte prognose.")

exante_table = doc.add_table(rows=1, cols=2)
exante_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(exante_table, ["✓ Tilgjengelig før produksjon", "✗ IKKE tilgjengelig (ekskludert)"], widths_in=[3.3, 3.3])
add_table_data_row(exante_table, ["Recoverable reserves (PDO-anslag)", "Faktisk kumulativ produksjon"], align_center=False)
add_table_data_row(exante_table, ["Antall planlagte brønner (PDO)", "Observert peak / decline / ramp"], align_center=False)
add_table_data_row(exante_table, ["Facility type (FPSO/Fixed/Subsea)", "Premium (krever 12 mnd data)"], align_center=False)
add_table_data_row(exante_table, ["API gravity (discovery DST)", "Realisert vannkutt"], align_center=False)
add_table_data_row(exante_table, ["Operatør (lisens-tildeling)", "Reservoaroverraskelser"], align_center=False)

add_heading("11.2 Peak-modell", level=2)
add_formula("ln(peak_MSm³/mnd) = −3.05 + 0.49·ln(recoverable) + 0.18·ln(n_wells) + facility")
add_para("Trent på 63 NCS-felt med kjent utfall. En viktig oppdagelse: forenkling SLO den kompliserte modellen. Vi testet 13 variabler (API, vanndybde, decade, operatør-dummies) men de fleste var ikke signifikante og skapte overfitting. Den enkle modellen (recoverable + brønner + facility) ga bedre out-of-sample-ytelse.")
add_para("Recoverable er den dominerende variabelen (β=0.49, p<0.001). Den konkave eksponenten (<1) betyr at peak ikke skalerer lineært med reserves — større felt får lengre platå, ikke proporsjonalt høyere peak, fordi facility-kapasitet er et tak.")

add_para("Duan smearing-korreksjon:", bold=True)
add_para("Når en modell fittes i log-rom og transformeres tilbake med exp(), får man medianen, ikke gjennomsnittet — en systematisk underprediksjon (Jensen's ulikhet). Vi korrigerer med Duan's smearing-estimator (multipliserer med gjennomsnittet av exp(residualer) = 1.07). Dette reduserer bias fra −17% til −11%.")

add_heading("11.3 Ramp, platå og decline sub-modeller", level=2)
submodel_table = doc.add_table(rows=1, cols=4)
submodel_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(submodel_table, ["Sub-modell", "Features", "LOO CV R²", "Out-of-sample"], widths_in=[1.6, 2.4, 1.3, 1.3])
add_table_data_row(submodel_table, ["Peak", "recov. + wells + facility", "0.84 (log)", "±35%"],
                   colors=[None, None, GREEN, ORANGE])
add_table_data_row(submodel_table, ["Ramp", "recoverable", "0.29", "±5 mnd"],
                   colors=[None, None, None, GREEN])
add_table_data_row(submodel_table, ["Plateau", "recoverable", "0.28", "±4 mnd"],
                   colors=[None, None, None, GREEN])
add_table_data_row(submodel_table, ["Decline", "recov. + wells + operatør", "0.35", "kalibrert"],
                   colors=[None, None, None, None])
add_caption("Tabell: Sub-modeller. Forenkling forbedret ramp (0.10→0.29) og platå (0.19→0.28) ved å fjerne overfitting-variabler.")

add_heading("11.4 Lifecycle-integrasjon med joint bootstrap", level=2)
add_para("De fire sub-modellene kombineres til én produksjonskurve (logistisk ramp + konstant platå + eksponentiell decline). To kritiske design-valg gjør usikkerheten realistisk:")
add_bullet("Joint bootstrap: samme felt-resampling brukes på tvers av alle fire modeller, slik at parameter-korrelasjoner bevares (store felt får både høy peak OG langt platå — ikke umulige kombinasjoner)")
add_bullet("Recovery-constraint: decline løses analytisk slik at kumulativ produksjon = recoverable reserves. Dette sikrer fysisk konsistens (recovery-ratio 1.00-1.01)")
add_para("Resultatet er P10/P50/P90-bånd som er fysisk plausible — ramp 21-29 mnd i stedet for det meningsløse 1-1626 mnd-båndet en naiv uavhengig sampling gir.")

# ═══════════════════════════════════════════════════════════════
# 12. HOLD-OUT VALIDERING (NY SEKSJON — V3)
# ═══════════════════════════════════════════════════════════════
add_heading("12. Out-of-sample hold-out validering", level=1)
add_para("Den viktigste ærlighetstesten: lat som et nylig felt er ukjent, re-tren modellen UTEN det, og forecast blindt fra kun PDO-input. Vi testet 13 felt som kom online nylig (Johan Sverdrup 2019, Edvard Grieg 2015, Goliat 2016, m.fl.).")

add_image("fig_holdout_validation.png", width_inches=6.5)
add_caption("Figur 6: Hold-out validering. Modellen så ALDRI testfeltet under trening. Ramp/platå treffer presist (±5/±4 mnd); peak-nivå har ~35% median feil — et range-verktøy, ikke punkt-estimat.")

add_heading("12.1 Det ærlige resultatet", level=2)
holdout_table = doc.add_table(rows=1, cols=3)
holdout_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(holdout_table, ["Komponent", "Out-of-sample resultat", "Dom"], widths_in=[2.0, 2.8, 1.8])
add_table_data_row(holdout_table, ["Ramp", "median feil ±5 mnd", "Punkt-estimat"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(holdout_table, ["Plateau", "median feil ±4 mnd", "Punkt-estimat"],
                   align_center=False, colors=[None, None, GREEN])
add_table_data_row(holdout_table, ["Peak (nivå)", "median feil ~35%", "Range-verktøy"],
                   align_center=False, colors=[None, None, ORANGE])
add_table_data_row(holdout_table, ["Decline (V5.1)", "nested CV R²=0.66", "Punkt-estimat"],
                   align_center=False, colors=[None, None, GREEN])

add_heading("12.2 Hvorfor 'log R²=0.84' og '35% feil' ikke motsier hverandre", level=2)
add_para("Peak-modellen har log-space CV R²=0.84 men lineær median feil ~35%. Dette er konsistent: log-R² belønner å treffe størrelsesorden, mens lineær %-feil måler absolutt presisjon. Et felt predikert 0.34 vs faktisk 0.53 er −37% lineært, men bare −0.44 i log-skala (liten på SD~1.0). For NPV er det den lineære feilen som teller — derfor rapporterer vi den ærlig.")

add_para("To regimer modellen ikke fanger med ex-ante data:", bold=True)
add_bullet("Mega-felt (Johan Sverdrup, −68%): ekstrem deliverabilitet krever reservoardata (permeabilitet, net pay) som per definisjon ikke er kjent før produksjon")
add_bullet("Små tie-back-felt (Maria, +102%): peak settes av facility-arrangement, ikke reservoaret")

add_heading("12.3 Korrekt bruk i ER", level=2)
add_para("Peak-modellen skal brukes til å triangulere mot operatør-guidance, ikke som et frittstående punkt-estimat. Dette er en sterkere ER-posisjon enn et oppblåst R²-tall — den tåler kritiske spørsmål i en modell-validering. Vi rapporterer det ekte out-of-sample-tallet (~35%), ikke det optimistiske log-R²=0.84 som om det betyr 16% feil.")

# ═══════════════════════════════════════════════════════════════
# 13. YGGDRASIL CASE STUDY (NY SEKSJON — V3)
# ═══════════════════════════════════════════════════════════════
add_heading("13. Case study: Yggdrasil (Aker BP NOAKA)", level=1)
add_para("Yggdrasil er Aker BPs store NOAKA-hub med planlagt produksjonsstart 2027. Vi anvender hele rammeverket på dette nye prosjektet — først på hub-nivå, deretter dekomponert i fem reservoarer.")

add_heading("13.1 Hub-nivå med triangulerte scenarier", level=2)
add_para("Siden recoverable reserves er den dominerende usikkerheten, kjører vi tre scenarier:")
ygg_table = doc.add_table(rows=1, cols=4)
ygg_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(ygg_table, ["Scenario", "Recoverable", "Peak P50", "vs guidance"], widths_in=[1.8, 1.8, 1.8, 1.8])
add_table_data_row(ygg_table, ["Low", "30 MSm³", "111 kboe/d", "1.3×"])
add_table_data_row(ygg_table, ["Base", "50 MSm³", "151 kboe/d", "1.8×"], bold=[True, True, True, True], bg=LIGHT_GREEN_BG)
add_table_data_row(ygg_table, ["High", "80 MSm³", "189 kboe/d", "2.2×"])
add_caption("Aker BP CMD-guidance: ~85 kboe/d olje peak.")

add_image("fig_yggdrasil_v2.png", width_inches=6.5)
add_caption("Figur 7: Yggdrasil V2 forecast. Øverst: hub-nivå tre scenarier. Midten/nederst: aggregat og fem komponenter (Krafla, Fulla, Frøy, Munin, Hugin) med staggered first oil og P10/P50/P90-bånd.")

add_heading("13.2 ER-tolkning", level=2)
add_para("Modellen sier konsistent over operatør-guidance (1.3-2.2×). Tre forklaringer, alle ER-relevante:")
add_bullet("PDO-konservatisme: operatører guider typisk lavt; NCS-felt outperformer ofte initial-guidance med 30-70%")
add_bullet("Vår modell underpredikerer systematisk (−11% bias selv etter smearing), så gapet til guidance kan være enda større")
add_bullet("Komponent-aggregatet (221 kboe/d) er høyere enn hub-nivå, da hver tie-back bidrar med egen peak")
add_para("Den ærlige ER-observasjonen: \"Yggdrasil viser oppside til Aker BP's produksjons-guidance, med et P10-P90-bånd på 119-189 kboe/d (hub-nivå, base scenario). Dette er en triangulering, ikke et presist punkt-estimat — peak-usikkerheten er ±35%.\"")

add_page_break()

# ═══════════════════════════════════════════════════════════════
# APPENDIX
# ═══════════════════════════════════════════════════════════════
add_heading("Appendix A: Modellutvikling-historikk", level=1)
dev_table = doc.add_table(rows=1, cols=4)
dev_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(dev_table, ["Trinn", "Forbedring", "CV R²", "Aker BP RMSE"],
                     widths_in=[0.8, 3.4, 1.2, 1.2])
add_table_data_row(dev_table, ["1", "Bare viskositet (Beggs-Robinson) med blend-API", "−0.015", "0.095"], align_center=False)
add_table_data_row(dev_table, ["2", "+ lifetime premium", "0.473", "0.076"], align_center=False)
add_table_data_row(dev_table, ["3", "Switch til siste 3 års premium", "0.544", "0.063"], align_center=False)
add_table_data_row(dev_table, ["4", "+ 12 mnd premium + |premium|  (V1 final)", "0.713", "0.056"], align_center=False)
add_table_data_row(dev_table, ["5", "+ Sodir DST + operatør research (110 felt)", "0.701", "0.064"], align_center=False)
add_table_data_row(dev_table, ["6", "+ Selektiv enrichment (V3, kun høy konfidens)", "0.713", "0.056"], align_center=False)
add_table_data_row(dev_table, ["7", "+ Felt-spesifikk T (V5.1 final) ★", "0.662*", "0.062"],
                   align_center=False, bold=[True, True, True, True],
                   colors=[None, None, GREEN, GREEN], bg=LIGHT_GREEN_BG)
add_caption("*Trinn 7 viser ærlig nested CV R²=0.662 (trinn 4-6 viser simple LOO som var litt optimistisk pga premium-sirkularitet, korrigert i V5.1).")

add_para("Pre-peak forecast-suite (V2) ble bygget på toppen av decline-modellen:")
fc_table = doc.add_table(rows=1, cols=3)
fc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(fc_table, ["Modell", "Out-of-sample", "Status"], widths_in=[2.6, 2.2, 2.0])
add_table_data_row(fc_table, ["Peak (med smearing)", "±35% median", "Range-verktøy"], align_center=False)
add_table_data_row(fc_table, ["Ramp duration", "±5 mnd", "Punkt-estimat"], align_center=False)
add_table_data_row(fc_table, ["Plateau duration", "±4 mnd", "Punkt-estimat"], align_center=False)
add_table_data_row(fc_table, ["Lifecycle (joint bootstrap)", "recovery 1.00-1.01", "Scenario-verktøy"], align_center=False)

add_heading("Appendix B: Variabel-definisjoner", level=1)
add_para("API gravity (reservoar): Tetthet av reservoarvæske målt på discovery well DST eller direct assay. NCS-typisk: 25° (tung) til 50° (lett kondensat).")
add_para("Viskositet (μ): Beggs-Robinson dead-oil viskositet i centipoise (cP), beregnet med felt-spesifikk T der tilgjengelig.")
add_para("D_annual: Eksponentiell decline rate på årlig basis. Production_year_T = Peak · exp(−D · T).")
add_para("Premium (P₁₂): Gjennomsnittlig log-avvik mellom faktisk og fysikk-forventet produksjon siste 12 måneder.")
add_para("Quality tier: operator_direct > dst_robust > operator_medium > dst_limited > operator_low > blend_inherited.")

add_heading("Appendix C: Master Fluid Library — Datakilder", level=1)
add_para("110 NCS-felt har innhentet reservoar-data fra:")
add_bullet("Sodir DST-database (sodir_wellbore_dst.csv): 1186 brønntester, 85 felt med oljedensitet")
add_bullet("Equinor crude assays: 15 høykonfidens med formasjon, reservoardybde, temperatur")
add_bullet("Aker BP CMD/research: 17 felt med API, reserver, eierandel")
add_bullet("Vår Energi research: Balder, Jotun, Ringhorne, Goliat, Marulk")
add_bullet("OKEA, Harbour Energy, DNO, ConocoPhillips: portfolio research")
add_bullet("Wikipedia, SPE-artikler, Norsk Petroleum, OffshoreTechnology.com")

add_para("Alle scripts og data er versjonskontrollert i prosjektet (analyses/decline_quality/).")

# ═══════════════════════════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════════════════════════
OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(str(OUTPUT))
print(f"Saved: {OUTPUT}")
print(f"Size: {OUTPUT.stat().st_size / 1024:.1f} KB")

# Copy to OneDrive
shutil.copy2(OUTPUT, ONEDRIVE)
print(f"Copied to OneDrive: {ONEDRIVE}")
